import json

from docx.shared import Inches
from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from docx import Document
import pyautogui
from plugin import Plugins

# 参考文档：https://www.justdopythonmain.py.com/2020/03/19/python-python-docs/

chrome_options = Options()
chrome_options.add_experimental_option("excludeSwitches", ['enable-automation'])
chrome_options.add_argument('ignore-certificate-errors')
driver = webdriver.Chrome(ChromeDriverManager().install(), options=chrome_options)
driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
    "source": """
        Object.defineProperty(navigator, 'webdriver', {
          get: () => undefined
        })
      """
})
#
#
# def gen_doc(query_word, spider_sites):
#     """
#         生成文档
#     """
#     num = 1
#     out = Document()
#     for site_name, site_url in spider_sites.items():
#         print("start query from site: %s" % site_url)
#         driver.get(site_url)
#         driver.maximize_window()
#         screen_shot_name = "screenshot/%s:%s.jpg" % (query_word, num)
#         pyautogui.screenshot(screen_shot_name)
#         out.add_paragraph("%s 查询结果:" % site_name)
#         out.add_picture(screen_shot_name, width=Inches(7.25))
#         num += 1
#     out.save("output/%s.doc" % query_word)
#


if __name__ == '__main__':
    query_word = "测试公司1"
    p = Plugins("spiders")
    plugins = p.load_plugins()
    for p in plugins:
        p_name = str(p.__name__)
        p_module_name = p_name.split('.')[1]
        print(p_module_name)
        if p_module_name == 'sse':
            print("start {}...".format(p_name))
            p.scrapy(query_word, driver)
            print("finished {}".format(p_name))
    driver.quit()
